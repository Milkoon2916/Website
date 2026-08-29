"""
문법 테스트 결과를 편집 가능한 .docx로 렌더링.
레퍼런스 이미지 스타일: 2단(newspaper column) 레이아웃, 문법 포인트 컬러 태그,
박스로 감싼 문제 문장, 원문자(①②③④⑤) 선택지.
나눔고딕을 기본 폰트로 지정 (한글 워드 문서 표준 폰트).
"""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

FONT_NAME = "나눔고딕"
ACCENT_RGB = RGBColor(0x3F, 0x46, 0x43)  # design-tokens.css --accent
TAG_BG_HEX = "EEEFEC"  # --accent-soft
BORDER_HEX = "C9C9C4"  # --border-strong

CIRCLED = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧"]


def _set_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 한글 폰트는 eastAsia 폰트를 별도로 지정해야 워드에서 정확히 적용됨
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _set_two_columns(section, num=2, space_dxa=560):
    """섹션을 신문처럼 num단으로 흐르게 설정 (python-docx엔 없는 기능이라 XML 직접 조작)."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(space_dxa))


def _add_tag(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"  {text}  ")
    _set_font(run, size=8.5, bold=True, color=RGBColor(0x6B, 0x6B, 0x66))
    # 태그 배경색 (paragraph shading)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), TAG_BG_HEX)
    pPr.append(shd)
    return p


def _add_boxed_paragraph(doc, text_runs, space_after=6):
    """text_runs: [(text, bold)] 리스트. 문제 문장을 사각 테두리 박스로 감싼 문단."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    for text, bold in text_runs:
        run = p.add_run(text)
        _set_font(run, size=10.5, bold=bold)
    # 사각 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '6')
        el.set(qn('w:color'), BORDER_HEX)
        pBdr.append(el)
    pPr.append(pBdr)
    # 안쪽 여백처럼 보이게 살짝 인덴트
    pPr_ind = OxmlElement('w:ind')
    pPr_ind.set(qn('w:left'), '80')
    pPr_ind.set(qn('w:right'), '80')
    pPr.append(pPr_ind)
    return p


def _add_plain(doc, text, size=10, bold=False, space_after=4, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    run.font.italic = italic
    return p


def _blank_line(doc, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BORDER_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(" ")
    _set_font(run, size=10)
    return p


def render_grammar_quiz_docx(result: dict, title: str = "문법 테스트") -> bytes:
    doc = Document()

    # 기본 스타일 폰트 지정
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ---------- 제목 (1단) ----------
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    _set_font(title_run, size=20, bold=True, color=ACCENT_RGB)
    title_p.paragraph_format.space_after = Pt(4)

    rule_p = doc.add_paragraph()
    rule_pPr = rule_p._p.get_or_add_pPr()
    rule_bdr = OxmlElement('w:pBdr')
    rule_bottom = OxmlElement('w:bottom')
    rule_bottom.set(qn('w:val'), 'single')
    rule_bottom.set(qn('w:sz'), '4')
    rule_bottom.set(qn('w:color'), BORDER_HEX)
    rule_bdr.append(rule_bottom)
    rule_pPr.append(rule_bdr)
    rule_p.paragraph_format.space_after = Pt(14)

    # 제목 아래부터 2단으로 전환 (연속 섹션 — 새 페이지로 넘어가지 않게)
    two_col_section = doc.add_section(WD_SECTION.CONTINUOUS)
    two_col_section.page_width = Cm(21.0)
    two_col_section.page_height = Cm(29.7)
    two_col_section.left_margin = Cm(1.8)
    two_col_section.right_margin = Cm(1.8)
    two_col_section.top_margin = Cm(1.0)
    two_col_section.bottom_margin = Cm(1.5)
    _set_two_columns(two_col_section, num=2)

    for q in result.get("questions", []):
        num = q.get("num")
        tag = q.get("tag", "")
        instruction = q.get("instruction", "")
        qtype = q.get("type")

        _add_tag(doc, tag)
        _add_plain(doc, f"{num}   {instruction}", size=10, bold=True, space_after=4)

        if qtype == "choice_parens" or qtype == "fill_blank_choice":
            _add_boxed_paragraph(doc, [(q.get("sentence", ""), False)])
            choices = q.get("choices", [])
            line = "     ".join(f"{CIRCLED[i]} {c}" for i, c in enumerate(choices))
            _add_plain(doc, line, size=10, space_after=14)

        elif qtype == "choose_sentence":
            choices = q.get("choices", [])
            lines = "\n".join(f"{CIRCLED[i]} {c}" for i, c in enumerate(choices))
            for i, c in enumerate(choices):
                _add_plain(doc, f"{CIRCLED[i]} {c}", size=10, space_after=1)
            _blank_line(doc, space_after=0)  # 여백용 (내용 없음)
            # 마지막 항목 뒤 여백 보정
            doc.paragraphs[-1].paragraph_format.space_after = Pt(14)

        elif qtype == "order_words":
            _add_plain(doc, q.get("korean_hint", ""), size=10, space_after=4, color=RGBColor(0x6B, 0x6B, 0x66))
            words_line = " / ".join(q.get("words", []))
            _add_boxed_paragraph(doc, [(words_line, False)])
            _blank_line(doc, space_after=14)

        elif qtype == "rewrite":
            _add_boxed_paragraph(doc, [(q.get("sentence", ""), False)])
            _blank_line(doc, space_after=14)

        else:
            # 알 수 없는 유형은 안전하게 문장만 표시
            if q.get("sentence"):
                _add_boxed_paragraph(doc, [(q.get("sentence", ""), False)])
            _blank_line(doc, space_after=14)

    # ---------- 정답 (마지막에 1단으로 별도 섹션) ----------
    doc.add_section()
    answer_section = doc.sections[-1]
    answer_section.page_width = Cm(21.0)
    answer_section.page_height = Cm(29.7)
    answer_section.left_margin = Cm(1.8)
    answer_section.right_margin = Cm(1.8)
    _set_two_columns(answer_section, num=1)

    _add_plain(doc, "정답", size=13, bold=True, space_after=8, color=ACCENT_RGB)
    for q in result.get("questions", []):
        num = q.get("num")
        qtype = q.get("type")
        if qtype in ("choice_parens", "fill_blank_choice", "choose_sentence"):
            idx = q.get("answer_index", 0)
            ans = CIRCLED[idx] if idx is not None and idx < len(CIRCLED) else "?"
        elif qtype in ("order_words", "rewrite"):
            ans = q.get("answer", "")
        else:
            ans = ""
        _add_plain(doc, f"{num}. {ans}", size=10, space_after=2)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
