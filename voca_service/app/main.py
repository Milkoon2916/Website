import json
import re
import time
import uuid
from html import escape as html_escape
from pathlib import Path

from fastapi import FastAPI, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from google import genai
from google.genai import types as genai_types
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = BASE_DIR / "app" / "templates"
STATIC_DIR = BASE_DIR / "static"
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

# PDF는 화면(result.html/worksheet.html)과 동일한 핑크 브랜드 스타일의 HTML을
# WeasyPrint로 그대로 인쇄하는 방식이다 (한글 폰트는 Dockerfile의 fonts-noto-cjk로 해결).
_pdf_env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))

DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
DEFAULT_GEMINI_MODEL = "gemini-2.5-flash"

app = FastAPI(title="ARA Vocab AI")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


@app.exception_handler(HTTPException)
def friendly_error_page(request: Request, exc: HTTPException):
    return HTMLResponse(
        status_code=exc.status_code,
        content=f"""<!DOCTYPE html>
<html lang="ko"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>오류 - ARA Vocab AI</title>
<link rel="stylesheet" href="static/style.css"></head>
<body><div class="container">
<header class="page-header">
  <div class="eyebrow">Error</div>
  <h1>문제가 발생했어요</h1>
  <p>{exc.detail}</p>
</header>
<a class="link-button" href="./">← 처음으로 돌아가기</a>
</div></body></html>""",
    )


# ---------------------------------------------------------------------------
# LLM 호출 (사용자가 입력한 provider / api_key / model 만 사용, 서버에 저장하지 않음)
# ---------------------------------------------------------------------------
def clean_json(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


TRANSIENT_MARKERS = ("503", "429", "UNAVAILABLE", "RESOURCE_EXHAUSTED", "overloaded", "rate_limit")


def _is_transient_error(exc: Exception) -> bool:
    text = str(exc)
    return any(marker in text for marker in TRANSIENT_MARKERS)


def call_llm(provider: str, api_key: str, model: str, system_prompt: str, user_prompt: str,
             max_retries: int = 4) -> str:
    """provider별로 사용자 API 키를 사용해 JSON 문자열 응답을 받아온다.
    api_key는 이 함수 호출 동안만 메모리에 존재하고 어디에도 저장되지 않는다.
    503(UNAVAILABLE)/429(RESOURCE_EXHAUSTED) 등 일시적인 과부하 에러는 지수 백오프로 자동 재시도한다."""
    if not api_key or not api_key.strip():
        raise HTTPException(status_code=400, detail="API 키를 입력해주세요.")

    provider = (provider or "gemini").strip().lower()
    last_error = None

    for attempt in range(max_retries):
        try:
            if provider == "openai":
                client = OpenAI(api_key=api_key.strip())
                response = client.chat.completions.create(
                    model=(model or DEFAULT_OPENAI_MODEL).strip(),
                    temperature=0.2,
                    response_format={"type": "json_object"},
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                )
                return response.choices[0].message.content

            if provider == "gemini":
                client = genai.Client(api_key=api_key.strip())
                response = client.models.generate_content(
                    model=(model or DEFAULT_GEMINI_MODEL).strip(),
                    contents=user_prompt,
                    config=genai_types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        response_mime_type="application/json",
                        temperature=0.2,
                    ),
                )
                return response.text

            raise HTTPException(status_code=400, detail=f"알 수 없는 provider입니다: {provider}")

        except HTTPException:
            raise
        except Exception as e:
            last_error = e
            if attempt < max_retries - 1 and _is_transient_error(e):
                # 1.5s, 3s, 6s, ... 지수 백오프 (모델 과부하/요청 한도 초과 시 자동 재시도)
                time.sleep(1.5 * (2 ** attempt))
                continue
            provider_label = "OpenAI" if provider == "openai" else "Gemini"
            raise HTTPException(status_code=400, detail=f"{provider_label} 요청 실패: {e}")

    provider_label = "OpenAI" if provider == "openai" else "Gemini"
    raise HTTPException(status_code=400, detail=f"{provider_label} 요청 실패(재시도 초과): {last_error}")


def analyze_passage(passage: str, level: str, count: int, focus: str,
                     provider: str, api_key: str, model: str) -> dict:
    prompt = f"""
You are an expert English teacher, vocabulary curriculum designer, and Korean EFL assessment writer.

Analyze the English passage below for {level} students.
Select exactly {count} high-value vocabulary items.

Primary focus: {focus}

Selection rules:
1. Prioritize words that are important for understanding the passage's topic, logic, argument, or narrative.
2. Prefer vocabulary with strong educational value and useful synonym/antonym expansion.
3. Avoid words that are too basic for the stated level unless the contextual meaning is especially important.
4. The meaning must reflect the meaning in this passage, not merely the first dictionary meaning.
5. The context_sentence must be copied exactly from the passage. Never invent an example sentence.
6. Give 2-4 useful synonyms and 1-3 useful antonyms when possible.
7. Include useful derivatives and collocations when appropriate.
8. Keep Korean explanations concise and classroom-friendly.
9. Rank importance from 1 to 5.
10. Difficulty should be one of: 초급, 중급, 고등, 수능.

Return ONLY valid JSON in exactly this structure:
{{
  "title": "short English title",
  "level": "{level}",
  "vocabulary": [
    {{
      "word": "string",
      "part_of_speech": "noun/verb/adjective/adverb/etc.",
      "meaning": "Korean meaning in context",
      "synonyms": ["string", "string"],
      "antonyms": ["string", "string"],
      "context_sentence": "exact sentence or sentence fragment copied from the passage",
      "context_meaning": "short Korean explanation",
      "importance": 1,
      "difficulty": "초급/중급/고등/수능",
      "derivatives": ["string"],
      "collocations": ["string"]
    }}
  ]
}}

Passage:
{passage}
"""
    raw = call_llm(provider, api_key, model, "Return valid JSON only. Do not add markdown.", prompt)
    return clean_json(raw)


def generate_worksheet(data: dict, worksheet_type: str, count: int, difficulty: str,
                        provider: str, api_key: str, model: str) -> dict:
    vocab_json = json.dumps(data.get("vocabulary", []), ensure_ascii=False)
    prompt = f"""
You are an expert Korean EFL test writer.

Create a vocabulary worksheet from the following vocabulary list.

Worksheet type: {worksheet_type}
Number of questions: {count}
Difficulty: {difficulty}

Question types may include:
- 영어 → 한국어 뜻
- 한국어 뜻 → 영어
- synonym multiple choice
- antonym multiple choice
- context cloze
- word order
- mixed

Important:
- Use only the vocabulary and information provided.
- Do not create impossible or ambiguous questions.
- For multiple choice, create exactly 4 choices.
- The answer must be unambiguous.
- Do not show the answer in the question text.

Return ONLY valid JSON:
{{
  "title": "Vocabulary Test",
  "questions": [
    {{
      "number": 1,
      "type": "meaning/synonym/antonym/cloze/order",
      "question": "question text",
      "answer": "correct answer",
      "choices": ["choice1", "choice2", "choice3", "choice4"]
    }}
  ]
}}

Vocabulary:
{vocab_json}
"""
    raw = call_llm(provider, api_key, model, "Return valid JSON only.", prompt)
    return clean_json(raw)


def generate_grammar_set(passage: str, target_grammar: str, level: str, set_number: int,
                          count: int, provider: str, api_key: str, model: str) -> dict:
    level_desc = "한국 중학교 수준 (중학 어법)" if level.startswith("중") else "한국 고등학교 수준 (고등/수능 어법)"
    prompt = f"""
You are an expert Korean EFL 어법(grammar) exam writer for {level_desc} students, in the style of
premium Korean grammar workbook publishers (e.g. category-tagged item banks like "중등발전" series).

Create SET {set_number} of a grammar practice test focused specifically on this target grammar point:
"{target_grammar}"

Number of questions in this set: exactly {count}. Every question must be different (no repeats or
near-duplicate sentences/patterns).

Break "{target_grammar}" down into varied sub-aspects and tag EACH question with the specific sub-point
it tests (field "tag"), e.g. if the target grammar is 관계대명사, sub-tags could include 주격 관계대명사,
목적격 관계대명사, 소유격 관계대명사, 관계대명사 what, 계속적 용법, 전치사+관계대명사, 관계부사와의 구별,
복합관계대명사, etc. Use as many different sub-tags across the set as make sense — do not reuse the same
narrow sub-tag more than 3-4 times in one set of {count}. If the target grammar is broader (e.g. "시제"),
invent similarly natural sub-aspects (예: 현재완료 vs 과거, 완료진행형, 시제 일치, 미래완료 등).

You may use the passage below as topical inspiration for some sentences, but you do not need to copy
its exact sentences — write original sentences whenever needed, as long as they clearly test the
target grammar point at a level appropriate for {level_desc} students.

Mix question TYPES naturally across the set — both "mc" (multiple choice) and "short" (서술형, no
choices, student writes the answer):
- mc, 괄호 안에서 알맞은 것 고르기: prompt has "(A / B)" inline, choices = ["A","B"] (2 choices).
- mc, 빈칸에 들어갈 가장 적절한 것/형태 고르기: prompt has a blank (write it as a long underscore run
  like "______"), choices = 4 options.
- mc, 두 빈칸에 들어갈 말로 가장 알맞은 것 고르기: prompt has two blanks, choices = 4 items each written
  as "first / second".
- mc, 어법상 옳은/틀린 문장 고르기: choices = 4 full candidate sentences.
- mc, 어법상 틀린 부분 고르기: prompt contains inline ① ~ ⑤ markers before 4-5 segments of one sentence,
  choices = the 4-5 underlined segment texts themselves (student picks which is wrong).
- short, 우리말과 같은 뜻이 되도록 주어진 단어를 올바른 순서로 배열하시오: prompt = the Korean sentence,
  prompt_secondary = the scrambled English words/phrases separated by " / ", choices = [] (empty),
  answer = the correctly ordered English sentence.
- short, 문장을 지시대로 바꿔쓰기 (예: 복합관계대명사를 이용해서 다시 쓰기, 수동태로 바꿔쓰기, 4형식으로
  바꿔쓰기 등): prompt = the original sentence + instruction context, choices = [], answer = the
  rewritten sentence.

Roughly 55-70% of the set should be "mc" and the rest "short", distributed naturally (do not group all
short-answer items together — interleave them).

Rules:
- Exactly one unambiguously correct answer per question (for "short" type, one clearly correct/expected
  answer, minor acceptable wording variants are fine but keep "answer" to the single best form).
- Provide a concise, classroom-friendly Korean explanation (해설) for every question, mc or short.
- Do not reveal the answer inside "instruction" or "prompt".
- Number questions 1 to {count} continuously.
- "prompt_secondary" is optional — only include it (non-empty) for word-ordering questions that need a
  separate scrambled word-bank box; otherwise omit it or set it to "".

Return ONLY valid JSON in exactly this structure:
{{
  "set_number": {set_number},
  "questions": [
    {{
      "number": 1,
      "tag": "short specific sub-grammar label in Korean, e.g. 주격 관계대명사",
      "type": "mc",
      "instruction": "괄호 안에서 알맞은 것을 고르시오.",
      "prompt": "the sentence/passage text for the question, with inline blanks/markers as needed",
      "prompt_secondary": "",
      "choices": ["choice1", "choice2"],
      "answer": "the correct choice or written answer, copied exactly if it's from choices",
      "explanation": "concise Korean explanation of why this is correct (and briefly why others are wrong, for mc)"
    }}
  ]
}}

Passage (topical context only, optional to use):
{passage[:4000]}
"""
    raw = call_llm(provider, api_key, model, "Return valid JSON only. Do not add markdown.", prompt)
    return clean_json(raw)

def generate_grammar_sets(passage: str, target_grammar: str, level: str, sets: int, count: int,
                           provider: str, api_key: str, model: str, title: str = "") -> dict:
    result_sets = []
    for i in range(1, sets + 1):
        result_sets.append(
            generate_grammar_set(passage, target_grammar, level, i, count, provider, api_key, model)
        )
    return {
        "title": title or "Grammar Practice",
        "target_grammar": target_grammar,
        "level": level,
        "sets": result_sets,
    }


# ---------------------------------------------------------------------------
# DOCX / PDF 생성 (기존 로직과 동일, 입력 데이터에만 의존)
# ---------------------------------------------------------------------------
def add_docx_header(doc, title, subtitle=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(subtitle).italic = True


def create_vocabulary_docx(data: dict, output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_docx_header(doc, "ARA VOCABULARY", f"{data.get('title', '')} · {data.get('level', '')}")

    for i, item in enumerate(data.get("vocabulary", []), 1):
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        rows = [
            ("단어", item.get("word", "")),
            ("품사", item.get("part_of_speech", "")),
            ("뜻", item.get("meaning", "")),
            ("유의어", ", ".join(item.get("synonyms", []))),
            ("반의어", ", ".join(item.get("antonyms", []))),
            ("지문 속 예문", item.get("context_sentence", "")),
            ("문맥상 의미", item.get("context_meaning", "")),
            ("파생어", ", ".join(item.get("derivatives", []))),
            ("연어/표현", ", ".join(item.get("collocations", []))),
        ]

        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        doc.add_paragraph()

    doc.save(output_path)


def create_worksheet_docx(data: dict, worksheet: dict, output_path: Path, answer_key=False):
    doc = Document()
    add_docx_header(doc, "ARA VOCABULARY TEST", data.get("title", ""))

    p = doc.add_paragraph("Name: ______________________________    Date: ________________")
    p.paragraph_format.space_after = Pt(14)

    for q in worksheet.get("questions", []):
        p = doc.add_paragraph()
        p.add_run(f"{q.get('number')}. ").bold = True

        if answer_key:
            p.add_run(f"{q.get('question', '')}  →  {q.get('answer', '')}")
        else:
            p.add_run(q.get("question", ""))
            choices = q.get("choices", [])
            if choices:
                doc.add_paragraph("    " + "   ".join(
                    f"{chr(65+i)}. {choice}" for i, choice in enumerate(choices)
                ))

        doc.add_paragraph("")

    doc.save(output_path)


def create_vocabulary_pdf(data: dict, output_path: Path):
    """화면에 보이는 word-card와 동일한 핑크 브랜드 스타일로 PDF를 렌더링한다."""
    template = _pdf_env.get_template("vocabulary_pdf.html.j2")
    html_str = template.render(
        title=data.get("title", ""),
        level=data.get("level", ""),
        vocabulary=data.get("vocabulary", []),
    )
    HTML(string=html_str).write_pdf(str(output_path))


def create_worksheet_pdf(data: dict, worksheet: dict, output_path: Path, answer_key=False):
    """화면 단어시험지와 동일한 스타일로 PDF를 렌더링한다. answer_key=True면 정답지 스타일."""
    template = _pdf_env.get_template("worksheet_pdf.html.j2")
    html_str = template.render(
        title=worksheet.get("title", data.get("title", "")),
        questions=worksheet.get("questions", []),
        answer_key=answer_key,
    )
    HTML(string=html_str).write_pdf(str(output_path))


def create_grammar_docx(grammar_data: dict, output_path: Path, answer_key=False):
    doc = Document()
    subtitle = f"목표 문법: {grammar_data.get('target_grammar', '')} · {grammar_data.get('level', '')}"
    add_docx_header(doc, "ARA GRAMMAR PRACTICE" + (" · 정답지" if answer_key else ""), subtitle)

    for s in grammar_data.get("sets", []):
        doc.add_paragraph("")
        h = doc.add_paragraph()
        h.add_run(f"SET {s.get('set_number')}").bold = True
        if not answer_key:
            p2 = doc.add_paragraph("Name: ______________________________    Date: ________________")
            p2.paragraph_format.space_after = Pt(10)

        for q in s.get("questions", []):
            tag = q.get("tag", "")
            if tag:
                pt = doc.add_paragraph()
                pt.add_run(f"[{tag}]").italic = True

            p = doc.add_paragraph()
            p.add_run(f"{q.get('number')}. ").bold = True
            p.add_run(q.get("instruction", q.get("question", "")))

            prompt = q.get("prompt", "")
            if prompt:
                doc.add_paragraph(f"    {prompt}")
            secondary = q.get("prompt_secondary", "")
            if secondary:
                doc.add_paragraph(f"    {secondary}")

            if answer_key:
                p2 = doc.add_paragraph(f"    → 정답: {q.get('answer', '')}")
                p2.add_run(f"\n    해설: {q.get('explanation', '')}")
            elif q.get("type") == "short":
                doc.add_paragraph("    답: ______________________________________________")
            else:
                choices = q.get("choices", [])
                if choices:
                    marks = ["①", "②", "③", "④", "⑤", "⑥"]
                    doc.add_paragraph("    " + "    ".join(
                        f"{marks[i] if i < len(marks) else i+1} {c}" for i, c in enumerate(choices)
                    ))
            doc.add_paragraph("")

    doc.save(output_path)


def create_grammar_pdf(grammar_data: dict, output_path: Path, answer_key=False):
    """세트별로 새 페이지에서 시작하는 어법 문제/정답지 PDF. 화면과 동일한 스타일을 그대로 인쇄한다."""
    template = _pdf_env.get_template("grammar_pdf.html.j2")
    html_str = template.render(
        title=grammar_data.get("title", ""),
        target_grammar=grammar_data.get("target_grammar", ""),
        level=grammar_data.get("level", ""),
        sets=grammar_data.get("sets", []),
        answer_key=answer_key,
    )
    HTML(string=html_str).write_pdf(str(output_path))


# ---------------------------------------------------------------------------
# 라우트
# ---------------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
def home():
    return (TEMPLATE_DIR / "index.html").read_text(encoding="utf-8")


GRAMMAR_SETS = 3
GRAMMAR_QUESTIONS_PER_SET = 30


@app.post("/generate-grammar-standalone", response_class=HTMLResponse)
def generate_grammar_standalone(
    provider: str = Form("gemini"),
    api_key: str = Form(""),
    model: str = Form(""),
    passage: str = Form(""),
    target_grammar: str = Form(""),
    grammar_level: str = Form("고등"),
    title: str = Form(""),
):
    """all-in-one 등 외부 페이지에서 iframe으로 바로 호출하는, 어법 문제 3세트 단독 생성 엔드포인트.
    /analyze처럼 어휘 분석을 같이 하지 않고 어법 문제만 생성해서 돌려준다."""
    if not passage.strip():
        raise HTTPException(status_code=400, detail="분석할 지문을 입력해주세요.")
    target_grammar = target_grammar.strip()
    if not target_grammar:
        raise HTTPException(status_code=400, detail="목표 어법을 입력해주세요.")

    grammar_data = generate_grammar_sets(
        passage, target_grammar, grammar_level, GRAMMAR_SETS, GRAMMAR_QUESTIONS_PER_SET,
        provider, api_key, model, title=title or "Grammar Practice",
    )
    section = render_grammar_section(grammar_data, target_grammar, grammar_level)

    return f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html_escape(target_grammar)} 어법 문제 - ARA Vocab AI</title>
<link rel="stylesheet" href="static/style.css">
</head>
<body>
<div class="container">
  {section}
</div>
</body>
</html>"""


@app.post("/analyze", response_class=HTMLResponse)
def analyze(
    provider: str = Form("gemini"),
    api_key: str = Form(""),
    model: str = Form(""),
    passage: str = Form(""),
    level: str = Form("고등학교 1학년"),
    count: int = Form(10),
    focus: str = Form("지문 이해와 시험 대비"),
    target_grammar: str = Form(""),
    grammar_level: str = Form("고등"),
):
    if not passage.strip():
        raise HTTPException(status_code=400, detail="분석할 지문을 입력해주세요.")
    data = analyze_passage(passage, level, count, focus, provider, api_key, model)
    data["passage"] = passage
    html = (TEMPLATE_DIR / "result.html").read_text(encoding="utf-8")

    cards = []
    for i, item in enumerate(data.get("vocabulary", []), 1):
        cards.append(f"""
        <article class="word-card">
            <div class="word-header">
                <div>
                    <span class="number">{i}</span>
                    <h2>{item.get('word', '')}</h2>
                </div>
                <span class="badge">{item.get('difficulty', '')}</span>
            </div>
            <div class="definition"><b>{item.get('meaning', '')}</b> · {item.get('part_of_speech', '')}</div>
            <div class="grid">
                <div><span>유의어</span><p>{', '.join(item.get('synonyms', []))}</p></div>
                <div><span>반의어</span><p>{', '.join(item.get('antonyms', []))}</p></div>
                <div><span>파생어</span><p>{', '.join(item.get('derivatives', []))}</p></div>
                <div><span>연어/표현</span><p>{', '.join(item.get('collocations', []))}</p></div>
            </div>
            <div class="sentence"><b>지문 속 예문</b><br>{item.get('context_sentence', '')}</div>
            <p class="context"><b>문맥상 의미:</b> {item.get('context_meaning', '')}</p>
        </article>
        """)

    grammar_section = ""
    target_grammar = target_grammar.strip()
    if target_grammar:
        try:
            grammar_data = generate_grammar_sets(
                passage, target_grammar, grammar_level, GRAMMAR_SETS, GRAMMAR_QUESTIONS_PER_SET,
                provider, api_key, model, title=data.get("title", ""),
            )
            grammar_section = render_grammar_section(grammar_data, target_grammar, grammar_level)
        except HTTPException as e:
            detail = e.detail if isinstance(e.detail, str) else "어법 문제 생성에 실패했습니다."
            grammar_section = f"""
  <header class="page-header" style="margin-top:48px;">
    <div class="eyebrow">Grammar Practice</div>
    <h1 style="font-size:26px;">목표 문법 어법 문제</h1>
    <p class="err" style="display:block;">어법 문제 생성에 실패했어요: {html_escape(detail)}<br>
      단어 분석 결과는 정상적으로 저장되었으니, 아래에서 다시 시도하려면 처음으로 돌아가 지문 분석을 다시 진행해주세요.</p>
  </header>
"""

    payload = json.dumps(data, ensure_ascii=False)
    html = html.replace("{{TITLE}}", data.get("title", "Vocabulary"))
    html = html.replace("{{LEVEL}}", data.get("level", ""))
    html = html.replace("{{VOCABULARY}}", "\n".join(cards))
    html = html.replace("{{DATA}}", payload)
    html = html.replace("{{PROVIDER}}", provider)
    html = html.replace("{{API_KEY}}", api_key)
    html = html.replace("{{MODEL}}", model)
    html = html.replace("{{GRAMMAR_SECTION}}", grammar_section)
    return html


def _gm_question_html(q: dict) -> str:
    marks = ["①", "②", "③", "④", "⑤", "⑥"]
    qtype = q.get("type", "mc")
    tag = q.get("tag", "")
    instruction = q.get("instruction", "")
    number = q.get("number", "")

    prompt_html = f"<div class='gm-prompt'>{q.get('prompt', '')}</div>" if q.get("prompt") else ""
    secondary = q.get("prompt_secondary", "")
    if secondary:
        prompt_html += f"<div class='gm-prompt gm-plain'>{secondary}</div>"

    body_html = ""
    if qtype == "short":
        body_html = "<div class='gm-writeline'></div>"
    else:
        choices = q.get("choices", [])
        if choices:
            choice_items = "".join(
                f"<span><span class='gm-num'>{marks[i] if i < len(marks) else i+1}</span>{c}</span>"
                for i, c in enumerate(choices)
            )
            grid_class = "gm-grid2" if len(choices) > 2 else ""
            body_html = f"<div class='gm-choices {grid_class}'>{choice_items}</div>"

    tag_html = f"<span class='gm-tag'>{tag}</span>" if tag else ""

    return f"""
    <div class="gm-qbox">
      {tag_html}
      <p class="gm-instruction"><b>{number}</b> {instruction}</p>
      {prompt_html}
      {body_html}
    </div>
    """


def render_grammar_section(grammar_data: dict, target_grammar: str, grammar_level: str) -> str:
    sets_html = []
    for s in grammar_data.get("sets", []):
        qs = [_gm_question_html(q) for q in s.get("questions", [])]
        n_q = len(s.get("questions", []))
        sets_html.append(f"""
        <div class="gm-set-title">SET {s.get('set_number')} ({n_q}문항)</div>
        <div class="gm-cols">
          {''.join(qs)}
        </div>
        """)

    grammar_json = html_escape(json.dumps(grammar_data, ensure_ascii=False), quote=True)

    return f"""
  <header class="page-header" style="margin-top:48px;">
    <div class="eyebrow">Grammar Practice</div>
    <h1 style="font-size:26px;">목표 문법 어법 문제 · {target_grammar}</h1>
    <p>{grammar_level} 수준으로 3세트, 세트당 30문항 이상(객관식+서술형 혼합)을 만들었어요. 문제지와 정답지를 각각 DOCX/PDF로 받을 수 있어요.</p>
  </header>

  <div class="toolbar">
    <form action="download/grammar/docx" method="post">
      <input type="hidden" name="data" value="{grammar_json}">
      <input type="hidden" name="answer_key" value="false">
      <button class="secondary" type="submit">📄 어법 문제 DOCX</button>
    </form>
    <form action="download/grammar/pdf" method="post">
      <input type="hidden" name="data" value="{grammar_json}">
      <input type="hidden" name="answer_key" value="false">
      <button class="secondary" type="submit">📕 어법 문제 PDF</button>
    </form>
    <form action="download/grammar/docx" method="post">
      <input type="hidden" name="data" value="{grammar_json}">
      <input type="hidden" name="answer_key" value="true">
      <button class="primary" type="submit">✅ 정답지 DOCX</button>
    </form>
    <form action="download/grammar/pdf" method="post">
      <input type="hidden" name="data" value="{grammar_json}">
      <input type="hidden" name="answer_key" value="true">
      <button class="primary" type="submit">✅ 정답지 PDF</button>
    </form>
  </div>

  <main>
    {''.join(sets_html)}
  </main>
"""


@app.post("/generate-worksheet", response_class=HTMLResponse)
def generate_worksheet_route(
    provider: str = Form("gemini"),
    api_key: str = Form(""),
    model: str = Form(""),
    data: str = Form(""),
    worksheet_type: str = Form("혼합형"),
    count: int = Form(20),
    difficulty: str = Form("중상"),
):
    if not data.strip():
        raise HTTPException(status_code=400, detail="단어 데이터가 없습니다. 지문 분석을 다시 진행해주세요.")
    parsed = json.loads(data)
    worksheet = generate_worksheet(parsed, worksheet_type, count, difficulty, provider, api_key, model)
    html = (TEMPLATE_DIR / "worksheet.html").read_text(encoding="utf-8")
    questions = []

    for q in worksheet.get("questions", []):
        choices = q.get("choices", [])
        choice_html = ""
        if choices:
            choice_items = "".join(
                f"<span>{chr(65+i)}. {c}</span>" for i, c in enumerate(choices)
            )
            choice_html = f"<div class='choices'>{choice_items}</div>"
        questions.append(
            f"<div class='question'><b>{q.get('number')}.</b> {q.get('question', '')}{choice_html}</div>"
        )

    html = html.replace("{{TITLE}}", worksheet.get("title", "Vocabulary Test"))
    html = html.replace("{{QUESTIONS}}", "\n".join(questions))
    html = html.replace("{{DATA}}", json.dumps(parsed, ensure_ascii=False))
    html = html.replace("{{WORKSHEET}}", json.dumps(worksheet, ensure_ascii=False))
    return html


@app.post("/download/vocabulary/docx")
def download_vocabulary_docx(data: str = Form(...)):
    parsed = json.loads(data)
    path = OUTPUT_DIR / f"ara_vocabulary_{uuid.uuid4().hex}.docx"
    create_vocabulary_docx(parsed, path)
    return FileResponse(path, filename="ara_vocabulary.docx",
                         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/download/vocabulary/pdf")
def download_vocabulary_pdf(data: str = Form(...)):
    parsed = json.loads(data)
    path = OUTPUT_DIR / f"ara_vocabulary_{uuid.uuid4().hex}.pdf"
    create_vocabulary_pdf(parsed, path)
    return FileResponse(path, filename="ara_vocabulary.pdf", media_type="application/pdf")


@app.post("/download/worksheet/docx")
def download_worksheet_docx(data: str = Form(...), worksheet: str = Form(...), answer_key: bool = Form(False)):
    parsed = json.loads(data)
    worksheet_data = json.loads(worksheet)
    path = OUTPUT_DIR / f"ara_worksheet_{uuid.uuid4().hex}.docx"
    create_worksheet_docx(parsed, worksheet_data, path, answer_key)
    filename = "ara_worksheet_answer_key.docx" if answer_key else "ara_worksheet.docx"
    return FileResponse(path, filename=filename,
                         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/download/worksheet/pdf")
def download_worksheet_pdf(data: str = Form(...), worksheet: str = Form(...), answer_key: bool = Form(False)):
    parsed = json.loads(data)
    worksheet_data = json.loads(worksheet)
    path = OUTPUT_DIR / f"ara_worksheet_{uuid.uuid4().hex}.pdf"
    create_worksheet_pdf(parsed, worksheet_data, path, answer_key)
    filename = "ara_worksheet_answer_key.pdf" if answer_key else "ara_worksheet.pdf"
    return FileResponse(path, filename=filename, media_type="application/pdf")


@app.post("/download/grammar/docx")
def download_grammar_docx(data: str = Form(...), answer_key: bool = Form(False)):
    grammar_data = json.loads(data)
    path = OUTPUT_DIR / f"ara_grammar_{uuid.uuid4().hex}.docx"
    create_grammar_docx(grammar_data, path, answer_key)
    filename = "ara_grammar_answer_key.docx" if answer_key else "ara_grammar.docx"
    return FileResponse(path, filename=filename,
                         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/download/grammar/pdf")
def download_grammar_pdf(data: str = Form(...), answer_key: bool = Form(False)):
    grammar_data = json.loads(data)
    path = OUTPUT_DIR / f"ara_grammar_{uuid.uuid4().hex}.pdf"
    create_grammar_pdf(grammar_data, path, answer_key)
    filename = "ara_grammar_answer_key.pdf" if answer_key else "ara_grammar.pdf"
    return FileResponse(path, filename=filename, media_type="application/pdf")
